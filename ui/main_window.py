# ui/main_window.py
import sys, os, fitz, re, json
from datetime import datetime

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QTextBrowser, QTextEdit, QLabel, QPushButton, QFileDialog,
    QSplitter, QToolBar, QLineEdit, QTreeWidgetItem, QTreeWidget,
    QTabWidget, QMessageBox, QButtonGroup, QSizePolicy, QComboBox,
    QScrollArea, QGridLayout, QToolButton, QDialog, QListWidget, QListWidgetItem
)
from PyQt6.QtGui import (
    QPixmap, QImage, QAction, QColor, QTextCursor, QTextCharFormat,
    QDesktopServices, QFont, QFontInfo, QFontMetrics, QTextTableFormat, QTextLength, QTextBlockFormat,
    QActionGroup, QIcon, QPainter, QKeySequence,
    QPainterPath, QShortcut, QTextDocument
)
from PyQt6.QtCore import Qt, QTimer, QUrl, QRect, QSize, QSizeF, QPointF, QRectF

# モジュール化したファイルをインポート
from ui.components import KanaInputFilter, ClickableLabel
from ui.dialogs import TimerSettingsDialog
from ui.widgets import PDFDisplayLabel, MemoWindow, AnswerSheet, TextAnnotationWidget, ShapeAnnotationWidget, ANSWER_TEMPLATE_PATH
from utils.constants import LAW_DATA
from utils.law_fetcher import LawFetcherThread
from utils.xml_parser import parse_law_xml_to_html, kanji_to_number_string


class MainWindow(QMainWindow):
    # (MainWindowクラスのコードをここに移動し、インポート文を修正)
    INITIAL_TIME = 120 * 60
    def __init__(self):
        super().__init__()
        self.setWindowTitle("司法試験等CBTシステム (シミュレーター版)")
        self.setGeometry(50, 50, 1600, 1000)
        self.current_answer_path = None
        self.memo_window = None
        self.layout_state = 0
        
        self.marker_palette = [
            ("黄", QColor("#fff176")),
            ("橙", QColor("#ffb74d")),
            ("緑", QColor("#aed581")),
            ("青", QColor("#64b5f6")),
            ("赤", QColor("#e57373")),
        ]
        self.pen_palette = [
            ("黒", QColor("black")),
            ("赤", QColor("#d32f2f")),
            ("青", QColor("#1976d2")),
        ]

        self.text_palette = [(label, QColor(color)) for label, color in self.marker_palette]
        self.text_sizes = {"small": 12, "medium": 16, "large": 20}
        self.text_size_key = "medium"
        self.text_color = QColor(self.text_palette[0][1])

        self.current_tool = "select"
        self.pen_color = QColor(self.pen_palette[0][1])
        self.pen_width = 2
        self.marker_color = QColor(self.marker_palette[0][1])
        self.marker_width = 10
        self.annotations = {}
        self.text_annotations = {}
        self.shape_annotations = {}
        self.selected_annotation = None
        self.pdf_scroll_area = None
        self.current_pdf_path = None
        self.page_overview_dialog = None
        self.spread_mode = False
        self.zoom_factor = 1.0
        self.MIN_ZOOM = 0.3
        self.MAX_ZOOM = 4.0
        self.fit_mode = None
        self.undo_stack = []
        self.redo_stack = []
        self.history_restoring = False
        self.max_history = 50
        self.law_bookmarks = []
        self._law_bookmark_dialog = None

        self.left_toolbar = self.create_left_toolbar()
        self.addToolBar(Qt.ToolBarArea.LeftToolBarArea, self.left_toolbar)
        
        self.problem_widget = self.create_problem_area()
        self.law_widget = self.create_law_area()
        self.answer_widget = self.create_answer_area()
        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.right_splitter = QSplitter(Qt.Orientation.Vertical)
        self._configure_splitter(self.main_splitter)
        self._configure_splitter(self.right_splitter)
        
        self.setup_toolbar_and_menu()
        self.reassemble_layout()
        self.connect_signals()
        self.setup_shortcuts()

        self.input_mode = 'roma'
        self.input_mode_filter = KanaInputFilter(self)
        QApplication.instance().installEventFilter(self.input_mode_filter)
        self.set_input_mode('roma')

        self.MIN_UI_FONT_SCALE = 0.5
        self.MAX_UI_FONT_SCALE = 1.5
        self.ui_font_scale = 1.0
        self._base_ui_font = QFont(QApplication.instance().font())
        self._base_timer_font = None
        self._base_law_font = None
        self._base_memo_font = None

        self.remaining_time = self.INITIAL_TIME
        self.timer = QTimer(self)
        self.timer.setInterval(1000)
        self.timer.timeout.connect(self.update_timer_display)
        self.update_timer_display(is_initial=True)
        self.timer.start()
        self.timer_paused = False
        self._closing_confirmed = False
        self._template_image_cache = {}
        
        self.law_fetcher_thread = None
        self._pending_page_render = None
        self._render_timer_active = False

    def createPopupMenu(self):
        return None

    def reassemble_layout(self):
        widgets = [self.problem_widget, self.law_widget, self.answer_widget]
        left_widget = widgets[self.layout_state % 3]
        top_right_widget = widgets[(self.layout_state + 1) % 3]
        bottom_right_widget = widgets[(self.layout_state + 2) % 3]

        previous_main_sizes = self.main_splitter.sizes() if self.main_splitter.count() else []
        previous_right_sizes = self.right_splitter.sizes() if self.right_splitter.count() else []

        left_widget.setParent(None)
        top_right_widget.setParent(None)
        bottom_right_widget.setParent(None)

        while self.right_splitter.count() > 0:
            self.right_splitter.widget(0).setParent(None)
        self.right_splitter.addWidget(top_right_widget)
        self.right_splitter.addWidget(bottom_right_widget)
        if previous_right_sizes and sum(previous_right_sizes) > 0:
            self.right_splitter.setSizes(previous_right_sizes)
        else:
            self.right_splitter.setSizes([400, 600])

        while self.main_splitter.count() > 0:
            self.main_splitter.widget(0).setParent(None)
        self.main_splitter.addWidget(left_widget)
        self.main_splitter.addWidget(self.right_splitter)
        if previous_main_sizes and sum(previous_main_sizes) > 0:
            self.main_splitter.setSizes(previous_main_sizes)
        else:
            self.main_splitter.setSizes([800, 800])

        self.setCentralWidget(self.main_splitter)

        if self.pdf_document:
            self._pending_page_render = self.current_page
            if not self._render_timer_active:
                self._render_timer_active = True
                QTimer.singleShot(0, self._retry_show_page)
        
    def swap_layout(self):
        self.layout_state = (self.layout_state + 1) % 3
        self.reassemble_layout()

    def setup_toolbar_and_menu(self):
        # ▼▼▼【バグ修正】ツールバーのエリアを明示的に指定 ▼▼▼
        toolbar = QToolBar("メインツールバー")
        self.addToolBar(Qt.ToolBarArea.TopToolBarArea, toolbar)
        toolbar.setStyleSheet("""
            QToolBar { spacing: 4px; }
            QPushButton, QToolButton { 
                background-color: #f0f0f0; 
                border: 1px solid #c0c0c0; 
                padding: 5px 10px;
                border-radius: 4px;
            }
            QPushButton:checked, QToolButton:checked { 
                background-color: #cde;
                border: 1px solid #9ac;
            }
            QPushButton#FinishButton {
                background-color: #007bff;
                color: white;
                font-weight: bold;
            }
        """)

        self.exam_type_label = QLabel("問題ファイルを開いてください...")
        toolbar.addWidget(self.exam_type_label)
        toolbar.addSeparator()

        self.memo_button = QPushButton("メモ")
        toolbar.addWidget(self.memo_button)

        self.problem_button = QPushButton("問題")
        self.law_button = QPushButton("法文")
        self.answer_button = QPushButton("答案")
        self.problem_button.setCheckable(True); self.problem_button.setChecked(True)
        self.law_button.setCheckable(True); self.law_button.setChecked(True)
        self.answer_button.setCheckable(True); self.answer_button.setChecked(True)
        
        self.display_group = QButtonGroup(self)
        self.display_group.setExclusive(False)
        self.display_group.addButton(self.problem_button)
        self.display_group.addButton(self.law_button)
        self.display_group.addButton(self.answer_button)
        
        toolbar.addWidget(self.problem_button)
        toolbar.addWidget(self.law_button)
        toolbar.addWidget(self.answer_button)

        self.swap_button = QPushButton("入替え")
        toolbar.addWidget(self.swap_button)
        toolbar.addSeparator()

        self.copy_button = QPushButton("コピー")
        self.cut_button = QPushButton("切取り")
        self.paste_button = QPushButton("貼付け")
        toolbar.addWidget(self.copy_button); toolbar.addWidget(self.cut_button); toolbar.addWidget(self.paste_button)
        toolbar.addSeparator()

        self.input_mode_roma_button = QPushButton("ローマ字入力")
        self.input_mode_kana_button = QPushButton("かな入力")
        self.input_mode_roma_button.setCheckable(True)
        self.input_mode_kana_button.setCheckable(True)
        self.input_mode_group = QButtonGroup(self)
        self.input_mode_group.setExclusive(True)
        self.input_mode_group.addButton(self.input_mode_roma_button)
        self.input_mode_group.addButton(self.input_mode_kana_button)
        toolbar.addWidget(self.input_mode_roma_button); toolbar.addWidget(self.input_mode_kana_button)
        toolbar.addSeparator()

        self.zoom_in_button = QPushButton("拡大 (+)")
        self.zoom_out_button = QPushButton("縮小 (-)")
        self.word_save_button = QPushButton("Word保存")
        toolbar.addWidget(self.zoom_in_button); toolbar.addWidget(self.zoom_out_button)
        toolbar.addWidget(self.word_save_button)
        
        spacer = QWidget(); spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        toolbar.addWidget(spacer)
        
        self.timer_label = ClickableLabel()
        self.timer_label.setStyleSheet("font-size: 16pt; color: red; font-weight: bold;")
        toolbar.addWidget(QLabel("残り時間："))
        toolbar.addWidget(self.timer_label)
        self._base_timer_font = QFont(self.timer_label.font())
        self.timer_label.setToolTip("クリックしてタイマーを設定 / 中断")
        self.timer_label.clicked.connect(self.open_timer_settings)

        self.finish_button = QPushButton("終了")
        self.finish_button.setObjectName("FinishButton")
        toolbar.addWidget(self.finish_button)

    def create_left_toolbar(self):
        # (このメソッドは変更なし)
        left_toolbar = QToolBar("PDFツール")
        left_toolbar.setOrientation(Qt.Orientation.Vertical)
        left_toolbar.setMovable(False)
        left_toolbar.setStyleSheet("QToolButton { padding: 8px 4px; font-size: 10pt; } QToolBar { spacing: 2px; }")
        left_toolbar.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextUnderIcon)

        self.select_action = QAction("選択", self); self.select_action.setCheckable(True)
        self.hand_action = QAction("ハンド", self); self.hand_action.setCheckable(True)
        self.marker_action = QAction("マーカー", self); self.marker_action.setCheckable(True)
        self.pen_action = QAction("ペン", self); self.pen_action.setCheckable(True)
        self.text_action = QAction("テキスト", self); self.text_action.setCheckable(True)
        
        self.tool_group = QActionGroup(self)
        self.tool_group.setExclusive(True)
        self.tool_group.addAction(self.select_action)
        self.tool_group.addAction(self.hand_action)
        self.tool_group.addAction(self.marker_action)
        self.tool_group.addAction(self.pen_action)
        self.tool_group.addAction(self.text_action)
        self.select_action.setChecked(True)

        left_toolbar.addAction(self.select_action)
        left_toolbar.addAction(self.hand_action)
        left_toolbar.addAction(self.marker_action)
        left_toolbar.addAction(self.pen_action)
        left_toolbar.addAction(self.text_action)

        self.circle_action = QAction("図形(〇)", self)
        self.triangle_action = QAction("図形(△)", self)
        self.cross_action = QAction("図形(✕)", self)
        self.toc_toolbar_action = QAction("目次", self)
        left_toolbar.addAction(self.circle_action)
        left_toolbar.addAction(self.triangle_action)
        left_toolbar.addAction(self.cross_action)
        left_toolbar.addSeparator()
        left_toolbar.addAction(self.toc_toolbar_action)
        left_toolbar.addSeparator()
        self.zoom_in_action = QAction("拡大", self)
        self.zoom_out_action = QAction("縮小", self)
        left_toolbar.addAction(self.zoom_in_action)
        left_toolbar.addAction(self.zoom_out_action)
        left_toolbar.addSeparator()
        self.fit_height_action = QAction("縦幅合わせ", self)
        self.fit_width_action = QAction("横幅合わせ", self)
        left_toolbar.addAction(self.fit_height_action)
        left_toolbar.addAction(self.fit_width_action)
        
        page_widget = QWidget()
        page_layout = QVBoxLayout(page_widget)
        page_layout.setContentsMargins(0,0,0,0); page_layout.setSpacing(0)
        self.page_num_input = QLineEdit(); self.page_num_input.setFixedWidth(50)
        self.page_num_input.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.page_label = QLabel("/ -"); self.page_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        page_layout.addWidget(self.page_num_input)
        page_layout.addWidget(self.page_label)

        left_toolbar.addSeparator()
        left_toolbar.addWidget(page_widget)
        
        self.goto_first_action = QAction("|<", self)
        self.prev_page_action = QAction("<", self)
        self.next_page_action = QAction(">", self)
        self.goto_last_action = QAction(">|", self)
        left_toolbar.addAction(self.goto_first_action)
        left_toolbar.addAction(self.prev_page_action)
        left_toolbar.addAction(self.next_page_action)
        left_toolbar.addAction(self.goto_last_action)
        left_toolbar.addSeparator()
        
        self.scroll_toggle_action = QAction("スクロール", self)
        self.scroll_toggle_action.setCheckable(True)
        self.scroll_toggle_action.setChecked(False)
        left_toolbar.addAction(self.scroll_toggle_action)

        self.spread_toggle_action = QAction("見開き", self)
        self.spread_toggle_action.setCheckable(True)
        self.spread_toggle_action.setChecked(False)
        left_toolbar.addAction(self.spread_toggle_action)
        left_toolbar.addSeparator()

        self.undo_toolbar_action = QAction("元に戻す", self)
        self.redo_toolbar_action = QAction("やり直し", self)
        self.undo_toolbar_action.setEnabled(False)
        self.redo_toolbar_action.setEnabled(False)
        left_toolbar.addAction(self.undo_toolbar_action)
        left_toolbar.addAction(self.redo_toolbar_action)

        return left_toolbar

    def setup_shortcuts(self):
        self.undo_shortcut = QShortcut(QKeySequence.StandardKey.Undo, self)
        self.undo_shortcut.activated.connect(self.undo)
        self.redo_shortcut = QShortcut(QKeySequence.StandardKey.Redo, self)
        self.redo_shortcut.activated.connect(self.redo)
        self.redo_shortcut_alt = QShortcut(QKeySequence("Ctrl+Shift+Z"), self)
        self.redo_shortcut_alt.activated.connect(self.redo)

    def create_problem_area(self):
        # (このメソッドは変更なし)
        pdf_area = QWidget()
        pdf_layout = QVBoxLayout(pdf_area)
        
        controls = QHBoxLayout()
        self.open_pdf_button = QPushButton("問題PDFを開く")
        self.save_pdf_button = QPushButton("書き込みを保存")
        
        controls.addWidget(self.open_pdf_button)
        controls.addWidget(self.save_pdf_button)
        controls.addStretch()

        self.pdf_display_label = PDFDisplayLabel(self)
        self.pdf_display_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.pdf_display_label.setText("「問題PDFを開く」")
        self.pdf_display_label.setCursor(Qt.CursorShape.ArrowCursor)

        self.text_tool_panel = QWidget()
        text_panel_layout = QHBoxLayout(self.text_tool_panel)
        text_panel_layout.setContentsMargins(0, 0, 0, 0)
        text_panel_layout.setSpacing(8)
        text_panel_layout.addWidget(QLabel("テキストサイズ:"))
        self.text_size_combo = QComboBox()
        self.text_size_combo.addItem("小", "small")
        self.text_size_combo.addItem("中", "medium")
        self.text_size_combo.addItem("大", "large")
        self.text_size_combo.setCurrentIndex(1)
        text_panel_layout.addWidget(self.text_size_combo)

        text_panel_layout.addWidget(QLabel("色:"))
        self.text_color_combo = QComboBox()
        for label, color in self.text_palette:
            self.text_color_combo.addItem(label, QColor(color))
            index = self.text_color_combo.count() - 1
            self.text_color_combo.setItemData(index, QColor(color), Qt.ItemDataRole.DecorationRole)
        self.text_color_combo.setCurrentIndex(0)
        text_panel_layout.addWidget(self.text_color_combo)
        text_panel_layout.addStretch()
        self.text_tool_panel.setVisible(False)

        self.pdf_scroll_area = QScrollArea()
        self.pdf_scroll_area.setWidgetResizable(False)
        self.pdf_scroll_area.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.pdf_scroll_area.setStyleSheet("QScrollArea { background: #ffffff; border: none; }")
        self.pdf_scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.pdf_scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.pdf_scroll_area.setWidget(self.pdf_display_label)

        pdf_layout.addLayout(controls)
        pdf_layout.addWidget(self.text_tool_panel)
        pdf_layout.addWidget(self.pdf_scroll_area, 1)

        self.pdf_document = None; self.current_page = 0; self.total_pages = 0
        self.update_text_style_from_controls()
        return pdf_area

    def create_law_area(self):
        # (このメソッドは変更なし)
        area = QWidget(); layout = QVBoxLayout(area)
        top_bar_layout = QHBoxLayout()
        self.toc_search_input = QLineEdit(); self.toc_search_input.setPlaceholderText("Q 目次検索...")
        self.toggle_toc_view_button = QPushButton("<"); self.toggle_toc_view_button.setCheckable(True); self.toggle_toc_view_button.setChecked(True); self.toggle_toc_view_button.setFixedWidth(30)
        self.toc_button = QPushButton("目次"); self.toc_button.setCheckable(True); self.toc_button.setChecked(True)
        jump_layout = QHBoxLayout(); jump_layout.addWidget(QLabel("第"))
        self.article_jump_input = QLineEdit(); self.article_jump_input.setFixedWidth(40)
        jump_layout.addWidget(self.article_jump_input); jump_layout.addWidget(QLabel("条"))
        self.paragraph_jump_input = QLineEdit(); self.paragraph_jump_input.setFixedWidth(30)
        jump_layout.addWidget(self.paragraph_jump_input); jump_layout.addWidget(QLabel("項"))
        self.item_jump_input = QLineEdit(); self.item_jump_input.setFixedWidth(30)
        jump_layout.addWidget(self.item_jump_input); jump_layout.addWidget(QLabel("号"))
        self.jump_button = QPushButton("移動")
        jump_layout.addWidget(self.jump_button)
        
        self.bookmark_button = QPushButton("付箋")
        self.law_search_toggle_button = QPushButton("検索"); self.law_search_toggle_button.setCheckable(True)
        top_bar_layout.addWidget(self.toc_search_input); top_bar_layout.addWidget(self.toggle_toc_view_button); top_bar_layout.addWidget(self.toc_button)
        top_bar_layout.addStretch(); top_bar_layout.addLayout(jump_layout); top_bar_layout.addWidget(self.bookmark_button); top_bar_layout.addWidget(self.law_search_toggle_button)
        
        self.law_search_bar = QWidget(); law_search_layout = QHBoxLayout(self.law_search_bar)
        self.law_search_input = QLineEdit(); self.law_search_input.setPlaceholderText("法文内を検索...")
        law_search_layout.addWidget(self.law_search_input)
        self.law_search_bar.setVisible(False)
        law_splitter = QSplitter(Qt.Orientation.Horizontal); self._configure_splitter(law_splitter)
        self.law_toc_tree = QTreeWidget(); self.law_toc_tree.setHeaderHidden(True); self.populate_law_tree()
        self.law_main_area = QTextBrowser(); self.law_main_area.setOpenExternalLinks(True)
        law_splitter.addWidget(self.law_toc_tree); law_splitter.addWidget(self.law_main_area); law_splitter.setSizes([300, 700])
        layout.addLayout(top_bar_layout); layout.addWidget(self.law_search_bar); layout.addWidget(law_splitter)
        return area

    def _configure_splitter(self, splitter):
        if not splitter:
            return
        splitter.setHandleWidth(8)
        splitter.setStyleSheet(
            """
            QSplitter::handle {
                background-color: #d0d8ec;
                border: 1px solid #7f91c8;
            }
            QSplitter::handle:hover {
                background-color: #b0bee6;
            }
            QSplitter::handle:pressed {
                background-color: #8297dd;
            }
            """
        )

    def populate_law_tree(self):
        # (このメソッドは変更なし)
        self.law_toc_tree.clear()
        for subject, laws in LAW_DATA.items():
            subject_item = QTreeWidgetItem(self.law_toc_tree); subject_item.setText(0, subject)
            for law_name, law_id in laws.items():
                law_item = QTreeWidgetItem(subject_item); law_item.setText(0, law_name); law_item.setData(0, Qt.ItemDataRole.UserRole, law_id)
        self.law_toc_tree.expandAll()

    def create_answer_area(self):
        # (このメソッドは変更なし)
        self.answer_tab_widget = QTabWidget()
        sheet1 = AnswerSheet(); sheet1.set_main_window(self)
        sheet2 = AnswerSheet(); sheet2.set_main_window(self)
        self.answer_tab_widget.addTab(sheet1, "第1問"); self.answer_tab_widget.addTab(sheet2, "第2問")
        return self.answer_tab_widget

    def connect_signals(self):
        # (このメソッドは変更なし)
        self.memo_button.clicked.connect(self.toggle_memo_window)
        self.bookmark_button.clicked.connect(self.handle_law_bookmark_action)
        self.zoom_in_button.clicked.connect(lambda: self.adjust_ui_font_scale(1.1))
        self.zoom_out_button.clicked.connect(lambda: self.adjust_ui_font_scale(1/1.1))
        self.word_save_button.clicked.connect(self.save_as_answer)
        self.input_mode_roma_button.clicked.connect(lambda checked: checked and self.set_input_mode('roma'))
        self.input_mode_kana_button.clicked.connect(lambda checked: checked and self.set_input_mode('kana'))
        self.problem_button.toggled.connect(self.problem_widget.setVisible)
        self.law_button.toggled.connect(self.law_widget.setVisible)
        self.answer_button.toggled.connect(self.answer_widget.setVisible)
        self.swap_button.clicked.connect(self.swap_layout)
        self.copy_button.clicked.connect(self.handle_copy)
        self.cut_button.clicked.connect(self.handle_cut)
        self.paste_button.clicked.connect(self.handle_paste)
        self.finish_button.clicked.connect(self.close)

        self.tool_group.triggered.connect(self.on_tool_selected)

        self.open_pdf_button.clicked.connect(self.open_pdf_file)
        self.save_pdf_button.clicked.connect(self.save_annotations_to_pdf)

        self.circle_action.triggered.connect(lambda: self.add_shape_annotation("circle"))
        self.triangle_action.triggered.connect(lambda: self.add_shape_annotation("triangle"))
        self.cross_action.triggered.connect(lambda: self.add_shape_annotation("cross"))
        self.toc_toolbar_action.triggered.connect(self.show_page_overview)

        self.text_size_combo.currentIndexChanged.connect(self.update_text_style_from_controls)
        self.text_color_combo.currentIndexChanged.connect(self.update_text_style_from_controls)

        self.prev_page_action.triggered.connect(self.show_prev_page)
        self.next_page_action.triggered.connect(self.show_next_page)
        self.goto_first_action.triggered.connect(lambda: self.show_page(0))
        self.goto_last_action.triggered.connect(lambda: self.show_page(self.total_pages - 1))
        self.page_num_input.returnPressed.connect(self.goto_page_from_input)
        self.zoom_in_action.triggered.connect(lambda: self.adjust_zoom(1.25))
        self.zoom_out_action.triggered.connect(lambda: self.adjust_zoom(0.8))
        self.scroll_toggle_action.toggled.connect(self.toggle_scroll_mode)
        self.spread_toggle_action.toggled.connect(self.toggle_spread_mode)
        self.undo_toolbar_action.triggered.connect(self.undo)
        self.redo_toolbar_action.triggered.connect(self.redo)
        self.fit_height_action.triggered.connect(self.fit_to_height)
        self.fit_width_action.triggered.connect(self.fit_to_width)

        self.answer_tab_widget.currentChanged.connect(lambda _: self.update_char_count())

        self.law_toc_tree.currentItemChanged.connect(self.on_law_tree_selection_changed)
        self.toc_search_input.textChanged.connect(self.filter_law_tree)
        self.toc_button.toggled.connect(self.toggle_law_toc_visibility); self.toggle_toc_view_button.toggled.connect(self.toggle_law_toc_visibility)
        self.jump_button.clicked.connect(self.jump_to_article); self.article_jump_input.returnPressed.connect(self.jump_to_article)
        self.law_search_toggle_button.toggled.connect(self.law_search_bar.setVisible)
        self.law_search_input.textChanged.connect(self.search_in_law_text)
        for i in range(self.answer_tab_widget.count()):
            sheet = self.answer_tab_widget.widget(i)
            sheet.contentChanged.connect(self.update_char_count)
            sheet.contentChanged.connect(self.register_snapshot)
            sheet.answer_undo_button.clicked.connect(sheet.undo_current)
            sheet.answer_redo_button.clicked.connect(sheet.redo_current)
            sheet.toggle_search_button.toggled.connect(sheet.search_replace_bar.setVisible)
            sheet.answer_search_input.textChanged.connect(lambda text, s=sheet: self.highlight_all_in_answer(s, text))
            sheet.answer_replace_button.clicked.connect(lambda _, s=sheet: self.replace_in_answer(s))
            sheet.answer_replace_all_button.clicked.connect(lambda _, s=sheet: self.replace_all_in_answer(s))
            sheet.search_prev_button.clicked.connect(lambda _, s=sheet: self.find_in_answer(s, forward=False))
            sheet.search_next_button.clicked.connect(lambda _, s=sheet: self.find_in_answer(s, forward=True))

    def on_tool_selected(self, action):
        if action == self.pen_action:
            self.current_tool = "pen"
        elif action == self.marker_action:
            self.current_tool = "marker"
        elif action == self.text_action:
            self.current_tool = "text"
        elif action == self.hand_action:
            self.current_tool = "hand"
        else:
            self.current_tool = "select"

        self.text_tool_panel.setVisible(self.current_tool == "text")

        if self.current_tool == "text":
            self.pdf_display_label.setCursor(Qt.CursorShape.IBeamCursor)
        elif self.current_tool in ("pen", "marker"):
            self.pdf_display_label.setCursor(Qt.CursorShape.CrossCursor)
        elif self.current_tool == "hand":
            self.pdf_display_label.setCursor(Qt.CursorShape.OpenHandCursor)
        else:
            self.pdf_display_label.setCursor(Qt.CursorShape.ArrowCursor)

        if self.current_tool != "hand":
            self.pdf_display_label.cancel_hand_drag()
    
    def set_tool_color(self, color):
        qcolor = QColor(color)
        if self.current_tool == "pen":
            self.pen_color = QColor(qcolor)
        elif self.current_tool == "marker":
            self.marker_color = QColor(qcolor)

    def set_tool_width(self, size_str):
        pen_sizes = {"small": 1, "medium": 2, "large": 4}
        marker_sizes = {"small": 8, "medium": 12, "large": 16}
        if self.current_tool == "pen":
            self.pen_width = pen_sizes.get(size_str, 2)
        elif self.current_tool == "marker":
            self.marker_width = marker_sizes.get(size_str, 12)

    def update_text_style_from_controls(self):
        if not hasattr(self, "text_size_combo"):
            return
        size_key = self.text_size_combo.currentData()
        if size_key in self.text_sizes:
            self.text_size_key = size_key

        color_data = self.text_color_combo.currentData()
        if isinstance(color_data, QColor):
            self.text_color = QColor(color_data)

        focus_widget = self.focusWidget()
        if isinstance(focus_widget, QTextEdit):
            possible_annotation = focus_widget.parent()
            if isinstance(possible_annotation, TextAnnotationWidget):
                possible_annotation.set_text_style(self.text_color, self.get_text_point_size())
                if not self.history_restoring:
                    self.register_snapshot()

    def get_text_point_size(self):
        return self.text_sizes.get(self.text_size_key, 16)

    def get_text_color(self):
        return QColor(self.text_color)

    def adjust_zoom(self, multiplier):
        new_zoom = self.zoom_factor * multiplier
        new_zoom = max(self.MIN_ZOOM, min(self.MAX_ZOOM, new_zoom))
        if abs(new_zoom - self.zoom_factor) < 0.01:
            return
        self.fit_mode = None
        self.zoom_factor = new_zoom
        if self.pdf_document:
            self.show_page(self.current_page)
            if not self.history_restoring:
                self.register_snapshot()

    def reset_zoom(self):
        self.zoom_factor = 1.0

    def fit_to_height(self):
        if not self.pdf_document:
            QMessageBox.warning(self, "縦幅合わせ", "PDFファイルが開かれていません。")
            return
        self.fit_mode = 'height'
        self.zoom_factor = 1.0
        self.show_page(self.current_page)
        if not self.history_restoring:
            self.register_snapshot()

    def fit_to_width(self):
        if not self.pdf_document:
            QMessageBox.warning(self, "横幅合わせ", "PDFファイルが開かれていません。")
            return
        self.fit_mode = 'width'
        self.zoom_factor = 1.0
        self.show_page(self.current_page)
        if not self.history_restoring:
            self.register_snapshot()

    def toggle_scroll_mode(self, horizontal_mode):
        if horizontal_mode and self.spread_mode:
            QMessageBox.information(self, "スクロール", "横スクロールでは見開き表示を利用できません。")
        self._apply_scroll_settings_without_refresh(horizontal_mode)
        if self.pdf_document:
            self.show_page(self.current_page)
            if not self.history_restoring:
                self.register_snapshot()

    def toggle_spread_mode(self, enabled):
        if enabled and self.scroll_toggle_action.isChecked():
            QMessageBox.information(self, "見開き", "横スクロール中は見開きを利用できません。")
            self.spread_toggle_action.blockSignals(True)
            self.spread_toggle_action.setChecked(False)
            self.spread_toggle_action.blockSignals(False)
            enabled = False
        self.spread_mode = enabled
        if self.spread_mode and self.pdf_document:
            # ensure current page visible remains within bounds
            if self.current_page >= self.total_pages:
                self.current_page = max(0, self.total_pages - 1)
        if self.pdf_document:
            self.show_page(self.current_page)
            if not self.history_restoring:
                self.register_snapshot()

    def clear_selection(self):
        if not self.selected_annotation:
            return
        selected = self.selected_annotation
        self.selected_annotation = None
        if hasattr(self, 'pdf_display_label') and self.pdf_display_label:
            try:
                self.pdf_display_label.cancel_selection_drag()
            except AttributeError:
                pass
        if selected.get('type') == 'text':
            widget = selected.get('widget')
            if widget:
                try:
                    widget.set_selected(False)
                except RuntimeError:
                    pass
        elif selected.get('type') == 'shape':
            widget = selected.get('widget')
            if widget:
                try:
                    widget.set_selected(False)
                    widget.clearFocus()
                except RuntimeError:
                    pass
        if hasattr(self, 'pdf_display_label') and self.pdf_display_label:
            self.pdf_display_label.update()

    def select_stroke_annotation(self, page, index):
        current = self.selected_annotation
        if current and current.get('type') == 'stroke' and current.get('page') == page and current.get('index') == index:
            return
        self.clear_selection()
        self.selected_annotation = {'type': 'stroke', 'page': page, 'index': index}
        if hasattr(self, 'pdf_display_label') and self.pdf_display_label:
            self.pdf_display_label.update()

    def select_text_annotation(self, widget):
        if not widget:
            return
        current = self.selected_annotation
        if current and current.get('type') == 'text' and current.get('widget') is widget:
            return
        self.clear_selection()
        self.selected_annotation = {'type': 'text', 'page': self.current_page, 'widget': widget}
        widget.set_selected(True)
        widget.raise_()

    def select_shape_annotation(self, widget):
        if not widget:
            return
        current = self.selected_annotation
        if current and current.get('type') == 'shape' and current.get('widget') is widget:
            return
        self.clear_selection()
        self.selected_annotation = {'type': 'shape', 'page': self.current_page, 'widget': widget}
        widget.set_selected(True)
        widget.setFocus()
        widget.raise_()

    def is_stroke_selected(self, page, index):
        selected = self.selected_annotation
        return (
            selected is not None and
            selected.get('type') == 'stroke' and
            selected.get('page') == page and
            selected.get('index') == index
        )

    def remove_text_annotation(self, page, widget):
        widgets = self.text_annotations.get(page)
        if not widgets:
            return
        if isinstance(widget, TextAnnotationWidget):
            if widget.text_edit.hasFocus():
                widget.text_edit.clearFocus()
            try:
                widget.set_selected(False)
            except RuntimeError:
                pass
        if self.selected_annotation and self.selected_annotation.get('widget') is widget:
            self.selected_annotation = None
        if widget in widgets:
            widgets.remove(widget)
            widget.deleteLater()
        if not widgets:
            self.text_annotations.pop(page, None)
        self.pdf_display_label.update()
        if not self.history_restoring:
            self.register_snapshot()

    def _clear_text_annotations(self):
        for widgets in self.text_annotations.values():
            for widget in widgets:
                if widget.text_edit.hasFocus():
                    widget.text_edit.clearFocus()
                try:
                    widget.set_selected(False)
                    widget.clearFocus()
                except RuntimeError:
                    pass
                widget.deleteLater()
        self.text_annotations.clear()
        if self.selected_annotation and self.selected_annotation.get('type') == 'text':
            self.selected_annotation = None
            self.pdf_display_label.update()

    def _update_text_annotations_visibility(self):
        for page, widgets in self.text_annotations.items():
            visible = page == self.current_page
            for widget in widgets:
                if not visible and widget.text_edit.hasFocus():
                    widget.text_edit.clearFocus()
                if not visible and self.selected_annotation and self.selected_annotation.get('widget') is widget:
                    widget.set_selected(False)
                    self.selected_annotation = None
                widget.setVisible(visible)
                if visible:
                    widget.raise_()
                    if self.selected_annotation and self.selected_annotation.get('widget') is widget:
                        widget.set_selected(True)

    def add_shape_annotation(self, shape_type):
        if not self.pdf_document:
            QMessageBox.warning(self, "図形ツール", "PDFファイルが開かれていません。")
            return
        self.pdf_display_label.add_shape_annotation(shape_type)

    def show_page_overview(self):
        if not self.pdf_document:
            QMessageBox.warning(self, "目次", "PDFファイルが開かれていません。")
            return

        if self.page_overview_dialog:
            self.page_overview_dialog.close()
            self.page_overview_dialog.deleteLater()
            self.page_overview_dialog = None

        dialog = QDialog(self)
        dialog.setWindowTitle("問題 目次")
        dialog.resize(960, 720)

        scroll_area = QScrollArea(dialog)
        scroll_area.setWidgetResizable(True)
        container = QWidget()
        grid_layout = QGridLayout(container)
        grid_layout.setContentsMargins(16, 16, 16, 16)
        grid_layout.setSpacing(18)

        columns = 3
        thumb_width = max(180, self.pdf_display_label.width() // (columns + 1)) if self.pdf_display_label.width() > 0 else 220
        for index in range(self.total_pages):
            thumb = self._render_page_thumbnail(index, thumb_width)

            button = QToolButton()
            button.setIcon(QIcon(thumb))
            button.setIconSize(thumb.size())
            button.setText(f"ページ {index + 1}")
            button.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextUnderIcon)
            button.setCursor(Qt.CursorShape.PointingHandCursor)
            button.setMinimumSize(thumb.size().width() + 20, thumb.size().height() + 60)
            button.clicked.connect(lambda _, p=index: self._handle_page_overview_click(p))

            row = index // columns
            col = index % columns
            grid_layout.addWidget(button, row, col)

        container.setLayout(grid_layout)
        scroll_area.setWidget(container)

        layout = QVBoxLayout(dialog)
        layout.addWidget(scroll_area)
        dialog.setLayout(layout)

        dialog.finished.connect(self._clear_page_overview_dialog)
        self.page_overview_dialog = dialog
        dialog.show()

    def remove_shape_annotation(self, page, widget):
        widgets = self.shape_annotations.get(page)
        if not widgets:
            return
        if self.selected_annotation and self.selected_annotation.get('widget') is widget:
            self.selected_annotation = None
        if widget in widgets:
            widgets.remove(widget)
            try:
                widget.set_selected(False)
                widget.clearFocus()
            except RuntimeError:
                pass
            widget.deleteLater()
        if not widgets:
            self.shape_annotations.pop(page, None)
        self.pdf_display_label.update()
        if not self.history_restoring:
            self.register_snapshot()

    def _clear_shape_annotations(self):
        for widgets in self.shape_annotations.values():
            for widget in widgets:
                try:
                    widget.set_selected(False)
                except RuntimeError:
                    pass
                widget.deleteLater()
        self.shape_annotations.clear()
        if self.selected_annotation and self.selected_annotation.get('type') == 'shape':
            self.selected_annotation = None
            self.pdf_display_label.update()

    def _update_shape_annotations_visibility(self):
        for page, widgets in self.shape_annotations.items():
            visible = page == self.current_page
            for widget in widgets:
                if visible:
                    widget.setVisible(True)
                    if self.selected_annotation and self.selected_annotation.get('widget') is widget:
                        widget.set_selected(True)
                    widget.raise_()
                else:
                    if widget.hasFocus():
                        widget.clearFocus()
                    widget.setVisible(False)
                    if self.selected_annotation and self.selected_annotation.get('widget') is widget:
                        widget.set_selected(False)
                        self.selected_annotation = None

    def _handle_page_overview_click(self, page_number):
        self.show_page(page_number)
        if self.page_overview_dialog:
            self.page_overview_dialog.close()

    def _clear_page_overview_dialog(self, *_):
        if self.page_overview_dialog:
            self.page_overview_dialog.deleteLater()
            self.page_overview_dialog = None

    def _render_page_image(self, page, scale, dpr):
        mat = fitz.Matrix(scale * dpr, scale * dpr)
        pix = page.get_pixmap(matrix=mat, annots=True)
        image_format = QImage.Format.Format_RGBA8888 if pix.alpha else QImage.Format.Format_RGB888
        image = QImage(pix.samples, pix.width, pix.height, pix.stride, image_format).copy()
        return image

    def _compose_spread_image(self, left_image, right_image, dpr, gap_px=20):
        if right_image is None:
            return left_image
        gap = int(gap_px * dpr)
        width = left_image.width() + right_image.width() + gap
        height = max(left_image.height(), right_image.height())
        combined = QImage(width, height, QImage.Format.Format_RGBA8888)
        combined.fill(Qt.GlobalColor.white)
        painter = QPainter(combined)
        painter.drawImage(0, 0, left_image)
        painter.drawImage(left_image.width() + gap, 0, right_image)
        painter.end()
        return combined

    def _render_page_thumbnail(self, page_number, target_width):
        page = self.pdf_document.load_page(page_number)
        rect = page.rect
        scale = target_width / rect.width if rect.width else 1.0
        scale = max(min(scale, 2.5), 0.1)
        mat = fitz.Matrix(scale, scale)
        pix = page.get_pixmap(matrix=mat, annots=False)

        image_format = QImage.Format.Format_RGBA8888 if pix.alpha else QImage.Format.Format_RGB888
        image = QImage(pix.samples, pix.width, pix.height, pix.stride, image_format).copy()
        pixmap = QPixmap.fromImage(image)
        return pixmap

    def save_annotations_to_pdf(self):
        current_sheet = self.answer_tab_widget.currentWidget()
        if not current_sheet:
            return

        default_dir = os.path.expanduser("~")
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        default_name = f"answer_{timestamp}.pdf"
        initial_path = os.path.join(default_dir, default_name)

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "PDF形式で保存",
            initial_path,
            "PDF Files (*.pdf)"
        )
        if not file_path:
            return
        if not file_path.lower().endswith('.pdf'):
            file_path += '.pdf'

        try:
            self._export_sheet_to_pdf(current_sheet, file_path)
        except Exception as exc:
            QMessageBox.critical(self, "保存エラー", f"PDFファイルの保存に失敗しました。\n{exc}")
            return

        QMessageBox.information(self, "保存完了", f"答案をPDF形式で保存しました。\n{file_path}")

    def register_snapshot(self):
        if self.history_restoring:
            return
        snapshot = self._serialize_state()
        if not snapshot:
            return
        if self.undo_stack and snapshot == self.undo_stack[-1]:
            return
        self.undo_stack.append(snapshot)
        if len(self.undo_stack) > self.max_history:
            self.undo_stack.pop(0)
        self.redo_stack.clear()
        self._update_history_actions()

    def undo(self):
        if len(self.undo_stack) <= 1:
            return
        current = self.undo_stack.pop()
        self.redo_stack.append(current)
        target = self.undo_stack[-1]
        self.history_restoring = True
        try:
            self._restore_state(target)
        finally:
            self.history_restoring = False
            self._update_history_actions()

    def redo(self):
        if not self.redo_stack:
            return
        target = self.redo_stack.pop()
        self.history_restoring = True
        try:
            self._restore_state(target)
        finally:
            self.history_restoring = False
        self.undo_stack.append(target)
        self._update_history_actions()

    def _update_history_actions(self):
        can_undo = len(self.undo_stack) > 1
        can_redo = len(self.redo_stack) > 0
        self.undo_toolbar_action.setEnabled(can_undo)
        self.redo_toolbar_action.setEnabled(can_redo)

    def _serialize_state(self):
        if not self.pdf_document:
            return {}
        h_val = self.pdf_scroll_area.horizontalScrollBar().value() if self.pdf_scroll_area else 0
        v_val = self.pdf_scroll_area.verticalScrollBar().value() if self.pdf_scroll_area else 0
        return {
            'annotations': self._serialize_annotations(),
            'text': self._serialize_text_annotations(),
            'shapes': self._serialize_shape_annotations(),
            'current_page': self.current_page,
            'zoom': self.zoom_factor,
            'spread': self.spread_mode,
            'fit_mode': self.fit_mode,
            'scroll_horizontal': self.scroll_toggle_action.isChecked() if self.pdf_scroll_area else False,
            'scroll_values': (h_val, v_val),
            'answers': self._serialize_answers()
        }

    def _restore_state(self, state):
        if not state:
            return
        self.history_restoring = True
        try:
            self.zoom_factor = state.get('zoom', 1.0)
            horizontal = state.get('scroll_horizontal', False)
            self.scroll_toggle_action.blockSignals(True)
            self.scroll_toggle_action.setChecked(horizontal)
            self.scroll_toggle_action.blockSignals(False)
            self._apply_scroll_settings_without_refresh(horizontal)

            spread = state.get('spread', False) and not horizontal
            self.spread_mode = spread
            self.spread_toggle_action.blockSignals(True)
            self.spread_toggle_action.setEnabled(not horizontal)
            self.spread_toggle_action.setChecked(spread)
            self.spread_toggle_action.blockSignals(False)

            fit_mode = state.get('fit_mode')
            if fit_mode not in ('width', 'height'):
                fit_mode = None
            if horizontal and fit_mode == 'width':
                fit_mode = None
            self.fit_mode = fit_mode

            self._deserialize_annotations(state.get('annotations', {}))
            self._deserialize_text_annotations(state.get('text', {}))
            self._deserialize_shape_annotations(state.get('shapes', {}))
            self._deserialize_answers(state.get('answers', []))
            target_page = max(0, min(state.get('current_page', self.current_page), self.total_pages - 1))
            self.show_page(target_page)
            scroll_values = state.get('scroll_values', (0, 0))
            if self.pdf_scroll_area:
                self.pdf_scroll_area.horizontalScrollBar().setValue(scroll_values[0])
                self.pdf_scroll_area.verticalScrollBar().setValue(scroll_values[1])
            self.clear_selection()
        finally:
            self.history_restoring = False

    def _apply_scroll_settings_without_refresh(self, horizontal_mode):
        if not self.pdf_scroll_area:
            return
        if horizontal_mode:
            self.pdf_scroll_area.setWidgetResizable(False)
            self.pdf_scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            self.pdf_scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            self.pdf_scroll_area.setAlignment(Qt.AlignmentFlag.AlignCenter)
            if self.spread_mode:
                self.spread_mode = False
                self.spread_toggle_action.blockSignals(True)
                self.spread_toggle_action.setChecked(False)
                self.spread_toggle_action.blockSignals(False)
            self.spread_toggle_action.setEnabled(False)
        else:
            self.pdf_scroll_area.setWidgetResizable(False)
            self.pdf_scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            self.pdf_scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            self.pdf_scroll_area.setAlignment(Qt.AlignmentFlag.AlignHCenter | Qt.AlignmentFlag.AlignTop)
            self.spread_toggle_action.setEnabled(True)

    def _serialize_annotations(self):
        data = {}
        for page, items in self.annotations.items():
            serialized_items = []
            for item in items:
                serialized_items.append({
                    'type': item['type'],
                    'color': item['color'].rgba(),
                    'width': item['width'],
                    'path': self._serialize_path(item['path'])
                })
            data[page] = serialized_items
        return data

    def _serialize_path(self, path):
        elements = []
        for i in range(path.elementCount()):
            elem = path.elementAt(i)
            elem_type_enum = getattr(elem, 'type', None)
            if elem_type_enum is None:
                elem_type = 0 if i == 0 else 1
            elif isinstance(elem_type_enum, int):
                elem_type = elem_type_enum
            else:
                elem_type = getattr(elem_type_enum, 'value', None)
                if elem_type is None:
                    # fallback using element position
                    elem_type = 0 if i == 0 else 1
            elements.append({'x': elem.x, 'y': elem.y, 'type': int(elem_type)})
        return elements

    def _deserialize_annotations(self, data):
        self.annotations = {}
        for page, items in data.items():
            page_int = int(page)
            restored = []
            for item in items:
                path = self._deserialize_path(item.get('path', []))
                restored.append({
                    'type': item.get('type', 'pen'),
                    'path': path,
                    'color': QColor.fromRgba(item.get('color', QColor('black').rgba())),
                    'width': item.get('width', 2)
                })
            if restored:
                self.annotations[page_int] = restored

    def _deserialize_path(self, elements):
        path = QPainterPath()
        for index, elem in enumerate(elements):
            point = QPointF(elem.get('x', 0.0), elem.get('y', 0.0))
            elem_type = elem.get('type', 0)
            if index == 0 or elem_type == 0:
                path.moveTo(point)
            else:
                path.lineTo(point)
        return path

    def _serialize_text_annotations(self):
        data = {}
        for page, widgets in self.text_annotations.items():
            serialized = []
            for widget in widgets:
                serialized.append({
                    'text': widget.text_edit.toPlainText(),
                    'x': widget.x(),
                    'y': widget.y(),
                    'width': widget.width(),
                    'height': widget.height(),
                    'color': widget._color.rgba(),
                    'font_point': widget._font_point
                })
            data[page] = serialized
        return data

    def _deserialize_text_annotations(self, data):
        self._clear_text_annotations()
        for page, items in data.items():
            page_int = int(page)
            restored_widgets = []
            for item in items:
                color = QColor.fromRgba(item.get('color', QColor('black').rgba()))
                font_point = item.get('font_point', 16)
                widget = TextAnnotationWidget(self.pdf_display_label, color, font_point)
                widget.delete_requested.connect(lambda w, p=page_int: self.pdf_display_label._handle_text_delete(p, w))
                widget.setGeometry(int(item.get('x', 0)), int(item.get('y', 0)), int(item.get('width', 220)), int(item.get('height', 80)))
                widget.text_edit.setPlainText(item.get('text', ''))
                widget.show()
                widget.raise_()
                restored_widgets.append(widget)
            if restored_widgets:
                self.text_annotations[page_int] = restored_widgets

    def _serialize_shape_annotations(self):
        data = {}
        for page, widgets in self.shape_annotations.items():
            serialized = []
            for widget in widgets:
                serialized.append({
                    'shape': widget.shape_type,
                    'x': widget.x(),
                    'y': widget.y(),
                    'width': widget.width(),
                    'height': widget.height()
                })
            data[page] = serialized
        return data

    def _deserialize_shape_annotations(self, data):
        self._clear_shape_annotations()
        for page, items in data.items():
            page_int = int(page)
            restored_widgets = []
            for item in items:
                size = QSize(int(item.get('width', 160)), int(item.get('height', 160)))
                widget = ShapeAnnotationWidget(self.pdf_display_label, item.get('shape', 'circle'), size)
                widget.delete_requested.connect(lambda w, p=page_int: self.pdf_display_label._handle_shape_delete(p, w))
                widget.move(int(item.get('x', 0)), int(item.get('y', 0)))
                widget.resize(size)
                widget.show()
                widget.raise_()
                restored_widgets.append(widget)
            if restored_widgets:
                self.shape_annotations[page_int] = restored_widgets

    def _serialize_answers(self):
        data = []
        for i in range(self.answer_tab_widget.count()):
            sheet = self.answer_tab_widget.widget(i)
            data.append({
                'page_texts': sheet.get_page_texts(),
                'current_page': sheet.current_page_index
            })
        return data

    def _deserialize_answers(self, data):
        if not data:
            return
        count = min(len(data), self.answer_tab_widget.count())
        for i in range(count):
            sheet = self.answer_tab_widget.widget(i)
            entry = data[i]
            texts = entry.get('page_texts', [])
            if texts:
                padded = texts + [''] * (sheet.TOTAL_PAGES - len(texts))
                sheet.set_page_texts(padded[:sheet.TOTAL_PAGES])
            sheet.set_current_page(entry.get('current_page', 0))
        self.update_char_count()

    def goto_page_from_input(self):
        try:
            page_num = int(self.page_num_input.text()) - 1
            if 0 <= page_num < self.total_pages:
                self.show_page(page_num)
        except ValueError:
            pass 

    def search_in_law_text(self, text):
        self.law_main_area.find(text)

    def handle_law_bookmark_action(self):
        cursor = self.law_main_area.textCursor()
        if cursor and cursor.hasSelection():
            self._add_law_bookmark(cursor)
        else:
            self._show_law_bookmark_dialog()

    def _add_law_bookmark(self, cursor):
        text = cursor.selectedText()
        if not text:
            QMessageBox.information(self, "付箋", "テキストを選択してから付箋ボタンを押してください。")
            return
        cleaned_text = text.replace('\u2029', '\n').strip()
        if not cleaned_text:
            QMessageBox.information(self, "付箋", "空白だけの選択は登録できません。")
            return

        start = cursor.selectionStart()
        end = cursor.selectionEnd()

        for bookmark in self.law_bookmarks:
            if bookmark['start'] == start and bookmark['end'] == end:
                QMessageBox.information(self, "付箋", "この選択範囲はすでに登録されています。")
                return

        snippet = cleaned_text if len(cleaned_text) <= 80 else cleaned_text[:77] + '…'
        self.law_bookmarks.append({
            'text': cleaned_text,
            'snippet': snippet,
            'start': start,
            'end': end,
        })
        QMessageBox.information(self, "付箋", "選択した条文を付箋として登録しました。")

    def _show_law_bookmark_dialog(self):
        if not self.law_bookmarks:
            QMessageBox.information(self, "付箋", "登録された付箋はありません。テキストを選択してから付箋ボタンを押してください。")
            return

        dialog = QDialog(self)
        dialog.setWindowTitle("法令付箋一覧")
        dialog.setModal(True)
        layout = QVBoxLayout(dialog)

        list_widget = QListWidget()
        for index, bookmark in enumerate(self.law_bookmarks, start=1):
            item = QListWidgetItem(f"{index}. {bookmark['snippet']}")
            list_widget.addItem(item)
        layout.addWidget(list_widget)

        button_row = QHBoxLayout()
        show_button = QPushButton("表示")
        delete_button = QPushButton("削除")
        close_button = QPushButton("閉じる")
        button_row.addWidget(show_button)
        button_row.addWidget(delete_button)
        button_row.addStretch()
        button_row.addWidget(close_button)
        layout.addLayout(button_row)

        def activate_selected():
            row = list_widget.currentRow()
            if row < 0:
                QMessageBox.information(dialog, "付箋", "表示する付箋を選択してください。")
                return
            self._activate_law_bookmark(row)
            dialog.accept()

        def delete_selected():
            row = list_widget.currentRow()
            if row < 0:
                QMessageBox.information(dialog, "付箋", "削除する付箋を選択してください。")
                return
            del self.law_bookmarks[row]
            list_widget.takeItem(row)
            if not self.law_bookmarks:
                dialog.accept()

        show_button.clicked.connect(activate_selected)
        delete_button.clicked.connect(delete_selected)
        close_button.clicked.connect(dialog.reject)
        list_widget.itemDoubleClicked.connect(lambda _: activate_selected())

        dialog.exec()

    def _activate_law_bookmark(self, index):
        if not (0 <= index < len(self.law_bookmarks)):
            return
        bookmark = self.law_bookmarks[index]
        cursor = self.law_main_area.textCursor()
        cursor.setPosition(bookmark['start'])
        cursor.setPosition(bookmark['end'], QTextCursor.MoveMode.KeepAnchor)
        self.law_main_area.setTextCursor(cursor)
        self.law_main_area.ensureCursorVisible()

    def pause_timer(self):
        if self.timer.isActive():
            self.timer.stop()
        self.timer_paused = True

    def resume_timer(self):
        if self.remaining_time <= 0:
            return
        if not self.timer.isActive():
            self.timer.start()
        self.timer_paused = False

    def open_timer_settings(self):
        dialog = TimerSettingsDialog(self, self.remaining_time, self.timer.isActive())
        if dialog.exec() == QDialog.DialogCode.Accepted:
            new_seconds = dialog.total_seconds()
            self.remaining_time = max(0, new_seconds)
            self.update_timer_display(is_initial=True)
            if self.remaining_time == 0:
                self.pause_timer()

    def prompt_exit(self):
        message = QMessageBox(self)
        message.setWindowTitle("終了確認")
        message.setIcon(QMessageBox.Icon.Question)
        message.setText("試験を終了しますか？")
        save_button = message.addButton("保存して終了", QMessageBox.ButtonRole.AcceptRole)
        exit_button = message.addButton("保存せず終了", QMessageBox.ButtonRole.DestructiveRole)
        cancel_button = message.addButton("キャンセル", QMessageBox.ButtonRole.RejectRole)
        message.exec()

        clicked = message.clickedButton()
        if clicked == save_button:
            return self.save_session_via_dialog()
        if clicked == exit_button:
            return True
        return False

    def set_input_mode(self, mode):
        if mode not in ('roma', 'kana'):
            return
        self.input_mode = mode
        roma = mode == 'roma'
        self.input_mode_roma_button.blockSignals(True)
        self.input_mode_kana_button.blockSignals(True)
        self.input_mode_roma_button.setChecked(roma)
        self.input_mode_kana_button.setChecked(not roma)
        self.input_mode_roma_button.blockSignals(False)
        self.input_mode_kana_button.blockSignals(False)
        self.input_mode_filter.set_mode('roma' if roma else 'kana')

    def adjust_ui_font_scale(self, multiplier):
        new_scale = self.ui_font_scale * multiplier
        new_scale = max(self.MIN_UI_FONT_SCALE, min(self.MAX_UI_FONT_SCALE, new_scale))
        self.apply_font_scale(new_scale)

    def apply_font_scale(self, scale):
        if abs(scale - getattr(self, 'ui_font_scale', 1.0)) < 1e-3:
            return
        self.ui_font_scale = scale
        app = QApplication.instance()
        base_font = getattr(self, '_base_ui_font', app.font())
        scaled_font = QFont(base_font)
        base_size = base_font.pointSizeF() if base_font.pointSizeF() > 0 else base_font.pointSize()
        if base_size > 0:
            scaled_font.setPointSizeF(base_size * scale)
        app.setFont(scaled_font)

        if hasattr(self, 'timer_label') and self._base_timer_font:
            timer_font = QFont(self._base_timer_font)
            base_size = timer_font.pointSizeF() if timer_font.pointSizeF() > 0 else timer_font.pointSize()
            if base_size > 0:
                timer_font.setPointSizeF(base_size * scale)
            self.timer_label.setFont(timer_font)

        self._apply_font_scale_to_law_area(scale)
        self._apply_font_scale_to_memo(scale)

    def _apply_font_scale_to_law_area(self, scale):
        if not hasattr(self, 'law_main_area'):
            return
        if self._base_law_font is None:
            self._base_law_font = QFont(self.law_main_area.font())
        base_font = self._base_law_font
        font = QFont(base_font)
        base_size = font.pointSizeF() if font.pointSizeF() > 0 else font.pointSize()
        if base_size > 0:
            font.setPointSizeF(base_size * scale)
        self.law_main_area.setFont(font)
        self.law_main_area.document().setDefaultFont(font)

    def _apply_font_scale_to_memo(self, scale):
        if self.memo_window:
            if not hasattr(self.memo_window, 'apply_font_scale'):
                return
            self.memo_window.apply_font_scale(scale)

    # --- Clipboard helpers -------------------------------------------------

    def _selected_text_annotation_editor(self):
        selected = getattr(self, 'selected_annotation', None)
        if selected and selected.get('type') == 'text':
            widget = selected.get('widget')
            text_edit = getattr(widget, 'text_edit', None)
            if text_edit and isinstance(text_edit, QTextEdit):
                return text_edit
        return None

    def _is_editable_widget(self, widget):
        if isinstance(widget, (QTextEdit, QPlainTextEdit)):
            return not widget.isReadOnly()
        if hasattr(widget, 'isReadOnly'):
            return not widget.isReadOnly()
        return False

    def _widget_has_selection(self, widget):
        if isinstance(widget, (QTextEdit, QPlainTextEdit)):
            cursor = widget.textCursor()
            return cursor is not None and cursor.hasSelection()
        if hasattr(widget, 'selectedText'):
            return bool(widget.selectedText())
        return False

    def _text_widget_candidates(self):
        candidates = []

        focus_widget = self.focusWidget()
        if isinstance(focus_widget, (QTextEdit, QPlainTextEdit)):
            candidates.append(focus_widget)

        annotation_editor = self._selected_text_annotation_editor()
        if annotation_editor:
            candidates.append(annotation_editor)

        if getattr(self, 'answer_tab_widget', None):
            current_sheet = self.answer_tab_widget.currentWidget()
            if current_sheet:
                editor = current_sheet.current_editor()
                if isinstance(editor, (QTextEdit, QPlainTextEdit)):
                    candidates.append(editor)

        if getattr(self, 'memo_window', None) and self.memo_window and self.memo_window.isVisible():
            memo_edit = getattr(self.memo_window, 'memo_edit', None)
            if isinstance(memo_edit, QTextEdit):
                candidates.append(memo_edit)

        return candidates

    def _resolve_clipboard_target(self, operation):
        candidates = self._text_widget_candidates()

        # For copy we also consider the law viewer even if it is read-only.
        if operation == 'copy' and getattr(self, 'law_main_area', None):
            if self.law_main_area not in candidates:
                candidates.append(self.law_main_area)

        for widget in candidates:
            if widget is None or not widget.isEnabled():
                continue
            if operation == 'copy':
                if self._widget_has_selection(widget):
                    return widget
            elif operation == 'cut':
                if self._is_editable_widget(widget) and self._widget_has_selection(widget):
                    return widget
            elif operation == 'paste':
                if self._is_editable_widget(widget):
                    return widget

        return None

    def handle_copy(self):
        widget = self._resolve_clipboard_target('copy')
        if widget and hasattr(widget, 'copy'):
            widget.setFocus(Qt.FocusReason.ShortcutFocusReason)
            widget.copy()

    def handle_cut(self):
        widget = self._resolve_clipboard_target('cut')
        if widget and hasattr(widget, 'cut'):
            widget.setFocus(Qt.FocusReason.ShortcutFocusReason)
            widget.cut()

    def handle_paste(self):
        widget = self._resolve_clipboard_target('paste')
        if widget and hasattr(widget, 'paste'):
            widget.setFocus(Qt.FocusReason.ShortcutFocusReason)
            widget.paste()

    def closeEvent(self, event):
        if getattr(self, '_closing_confirmed', False):
            event.accept()
            return
        if self.prompt_exit():
            self._closing_confirmed = True
            event.accept()
        else:
            event.ignore()

    def on_law_tree_selection_changed(self, current, previous):
        if current is None: return
        law_id = current.data(0, Qt.ItemDataRole.UserRole)
        if law_id: self.fetch_law_data(current.text(0), law_id)

    def filter_law_tree(self, text):
        for i in range(self.law_toc_tree.topLevelItemCount()):
            subject_item = self.law_toc_tree.topLevelItem(i); has_visible_child = False
            for j in range(subject_item.childCount()):
                law_item = subject_item.child(j); is_match = text.lower() in law_item.text(0).lower(); law_item.setHidden(not is_match)
                if is_match: has_visible_child = True
            subject_item.setHidden(not has_visible_child)
            if has_visible_child: subject_item.setExpanded(True)

    def on_law_data_ready(self, result_tuple):
        xml_string = result_tuple[0]
        _, main_html = parse_law_xml_to_html(xml_string)
        self.law_main_area.setHtml(main_html)

    def on_law_data_error(self, error_message): self.law_main_area.setText(error_message)

    def toggle_law_toc_visibility(self, checked):
        self.law_toc_tree.setVisible(checked); self.toc_search_input.setVisible(checked); self.toc_button.setChecked(checked); self.toggle_toc_view_button.setChecked(checked)
        self.toggle_toc_view_button.setText("<" if checked else ">")

    def jump_to_article(self):
        article_num = self.article_jump_input.text().strip()
        paragraph_num = self.paragraph_jump_input.text().strip()
        item_num = self.item_jump_input.text().strip()
        if not article_num: return
        
        anchor_name = f"article-{article_num}"
        if paragraph_num:
            anchor_name += f"-{paragraph_num}"
            if item_num:
                anchor_name += f"-{item_num}"
        self.law_main_area.scrollToAnchor(anchor_name)
    
    def fetch_law_data(self, law_name, law_id):
        if not law_id: self.law_main_area.setText("法令を選択してください。"); return
        self.law_main_area.setText(f"{law_name}のデータを取得中...");
        if self.law_fetcher_thread and self.law_fetcher_thread.isRunning(): self.law_fetcher_thread.terminate(); self.law_fetcher_thread.wait()
        self.law_fetcher_thread = LawFetcherThread(law_id, self)
        self.law_fetcher_thread.result_ready.connect(self.on_law_data_ready); self.law_fetcher_thread.error_occurred.connect(self.on_law_data_error); self.law_fetcher_thread.start()

    def save_as_answer(self):
        current_sheet = self.answer_tab_widget.currentWidget()
        if not current_sheet:
            return

        default_dir = os.path.expanduser("~")
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        default_name = f"answer_{timestamp}.docx"
        initial_path = os.path.join(default_dir, default_name)

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Word形式で保存",
            initial_path,
            "Word Documents (*.docx)"
        )
        if not file_path:
            return
        if not file_path.lower().endswith('.docx'):
            file_path += '.docx'

        try:
            self._export_sheet_to_docx(current_sheet, file_path)
        except Exception as exc:
            QMessageBox.critical(self, "保存エラー", f"Wordファイルの保存に失敗しました。\n{exc}")
            return

        self.current_answer_path = file_path
        QMessageBox.information(self, "保存完了", f"答案をWord形式で保存しました。\n{file_path}")

    def _export_sheet_to_docx(self, sheet, file_path):
        page_texts = self._sheet_page_texts(sheet)
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

        max_chars = sheet.pages[0].max_chars if sheet.pages else 30
        content_width_cm = section.page_width.cm - (section.left_margin.cm + section.right_margin.cm)
        line_col_width_cm = 1.0
        text_width_cm = max(content_width_cm - line_col_width_cm, 1.0)
        char_width_cm = text_width_cm / max_chars
        font_size_pt = max(8.0, min(char_width_cm * 72.0 / 2.54, 20.0))

        normal_style = doc.styles['Normal']
        normal_style.font.name = 'MS Mincho'
        normal_style.font.size = Pt(font_size_pt)
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')

        total_pages = len(page_texts)
        for page_index, page_text in enumerate(page_texts):
            lines = self._page_lines_for_export(page_text)

            table = doc.add_table(rows=23, cols=2)
            table.style = 'Table Grid'
            table.autofit = False
            line_col = table.columns[0]
            text_col = table.columns[1]
            line_width = Cm(1.0)
            line_col.width = line_width
            text_col.width = section.page_width - section.left_margin - section.right_margin - line_width

            for row_idx in range(23):
                row = table.rows[row_idx]
                line_cell = row.cells[0]
                text_cell = row.cells[1]
                line_cell.width = line_width
                text_cell.width = text_col.width

                line_cell.text = ''
                text_cell.text = ''

                line_para = line_cell.paragraphs[0]
                line_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                line_para.paragraph_format.space_after = Pt(0)
                line_para.paragraph_format.space_before = Pt(0)
                line_run = line_para.add_run(f"{row_idx + 1:02}")
                line_run.font.size = Pt(12)
                line_run.font.name = 'MS Mincho'
                line_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')

                text_para = text_cell.paragraphs[0]
                text_para.paragraph_format.space_after = Pt(0)
                text_para.paragraph_format.space_before = Pt(0)
                text_para.paragraph_format.line_spacing = 1
                line_text = lines[row_idx] if row_idx < len(lines) else ''
                text_run = text_para.add_run(line_text)
                text_run.font.size = Pt(12)
                text_run.font.name = 'MS Mincho'
                text_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')

            label_para = doc.add_paragraph(f"{page_index + 1}ページ目")
            label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label_para.paragraph_format.space_before = Pt(4)
            label_para.paragraph_format.space_after = Pt(0)
            label_run = label_para.runs[0]
            label_run.font.size = Pt(font_size_pt)
            label_run.font.name = 'MS Mincho'
            label_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')

            if page_index != total_pages - 1:
                doc.add_page_break()

        doc.save(file_path)

    def _export_sheet_to_pdf(self, sheet, file_path):
        page_texts = self._sheet_page_texts(sheet)
        if not page_texts:
            page_texts = ['']

        font = QFont("Hiragino Mincho ProN", 12)
        if not QFontInfo(font).exactMatch():
            font = QFont("MS Mincho", 12)

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(file_path)
        printer.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
        try:
            printer.setPageMargins(QMarginsF(20, 20, 20, 25), QPageLayout.Unit.Millimeter)
        except TypeError:
            try:
                printer.setPageMargins(20, 20, 20, 25, QPrinter.Unit.Millimeter)
            except Exception:
                pass

        painter = QPainter(printer)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        painter.setFont(font)
        metrics = QFontMetrics(font)

        max_lines = sheet.pages[0].max_lines if sheet.pages else 23
        max_chars = sheet.pages[0].max_chars if sheet.pages else 30
        char_width = metrics.horizontalAdvance('あ')
        line_height = metrics.lineSpacing()

        line_number_width = metrics.horizontalAdvance('00') + char_width * 0.4
        column_spacing = char_width * 0.5
        text_area_width = char_width * max_chars

        page_rect = printer.pageRect(QPrinter.Unit.Point)
        page_width = page_rect.width()
        page_height = page_rect.height()
        total_text_height = line_height * max_lines
        total_width = line_number_width + column_spacing + text_area_width
        start_x = page_rect.x() + (page_width - total_width) / 2.0
        top_margin = page_rect.y() + (page_height - total_text_height) / 2.0

        template_dpi = printer.resolution()
        template_dpi = max(72, min(template_dpi, 300))
        template_image = self._get_template_image(template_dpi)
        template_scale_x = template_scale_y = 1.0
        if not template_image.isNull():
            template_scale_x = page_rect.width() / template_image.width()
            template_scale_y = page_rect.height() / template_image.height()

        for page_index, page_text in enumerate(page_texts):
            if page_index > 0:
                printer.newPage()

            if not template_image.isNull():
                painter.save()
                painter.translate(page_rect.topLeft())
                painter.scale(template_scale_x, template_scale_y)
                painter.drawImage(QPointF(0, 0), template_image)
                painter.restore()

            lines = self._page_lines_for_export(page_text)
            for row_idx in range(max_lines):
                y = top_margin + row_idx * line_height
                line_rect = QRectF(start_x, y, line_number_width, line_height)
                painter.drawText(line_rect, int(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter), f"{row_idx + 1:02}")

                text_rect = QRectF(start_x + line_number_width + column_spacing, y, text_area_width, line_height)
                text = lines[row_idx] if row_idx < len(lines) else ''
                painter.drawText(text_rect, int(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter), text)

            footer_rect = QRectF(page_rect.x(), page_rect.bottom() - line_height * 1.5, page_rect.width(), line_height)
            painter.drawText(footer_rect, int(Qt.AlignmentFlag.AlignCenter), f"{page_index + 1} / {len(page_texts)}")

        painter.end()

    def _sheet_page_texts(self, sheet):
        texts = sheet.get_page_texts()
        last_nonempty = -1
        for idx, text in enumerate(texts):
            if text and text.strip():
                last_nonempty = idx
        if last_nonempty == -1:
            return [texts[0]] if texts else ['']
        return texts[:last_nonempty + 1]

    def _page_lines_for_export(self, text, max_lines=23, max_chars=30):
        lines = []
        raw_lines = text.split('\n') if text else ['']
        for raw in raw_lines:
            segment = raw
            if segment == '':
                lines.append('')
            else:
                while segment:
                    lines.append(segment[:max_chars])
                    segment = segment[max_chars:]
        lines = lines[:max_lines]
        while len(lines) < max_lines:
            lines.append('')
        return lines

    def _get_template_image(self, dpi):
        cache_key = int(dpi)
        if cache_key in self._template_image_cache:
            return self._template_image_cache[cache_key]
        image = QImage()
        if ANSWER_TEMPLATE_PATH and os.path.exists(ANSWER_TEMPLATE_PATH):
            try:
                doc = fitz.open(ANSWER_TEMPLATE_PATH)
                page = doc.load_page(0)
                scale = dpi / 72.0
                pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
                img_bytes = pix.tobytes('png')
                image.loadFromData(img_bytes, 'PNG')
                doc.close()
            except Exception:
                image = QImage()
        self._template_image_cache[cache_key] = image
        return image

    def _serialize_full_session(self):
        payload = {
            'version': 1,
            'saved_at': datetime.now().isoformat(),
            'remaining_time': self.remaining_time,
            'timer_paused': not self.timer.isActive() and self.remaining_time > 0,
            'input_mode': self.input_mode,
            'answers': self._serialize_answers(),
        }
        if self.current_pdf_path:
            payload['pdf'] = {
                'path': self.current_pdf_path,
                'state': self._serialize_state()
            }
        else:
            payload['pdf'] = None
        return payload

    def save_session_via_dialog(self):
        default_dir = os.path.expanduser("~")
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        default_name = f"cbt_session_{timestamp}.json"
        initial_path = os.path.join(default_dir, default_name)
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "セッションの保存先を選択",
            initial_path,
            "CBTセッション (*.json);;All Files (*)"
        )
        if not file_path:
            return False
        if not os.path.splitext(file_path)[1]:
            file_path += '.json'

        data = self._serialize_full_session()
        directory = os.path.dirname(file_path)
        try:
            if directory and not os.path.exists(directory):
                os.makedirs(directory, exist_ok=True)
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as exc:
            QMessageBox.critical(self, "保存エラー", f"セッションの保存に失敗しました。\n{exc}")
            return False


        QMessageBox.information(self, "保存完了", f"セッションを保存しました。\n{file_path}")
        return True

    def show_ai_review_placeholder(self):
        QMessageBox.information(self, "AI添削（開発中）", "AIによる答案添削機能は現在開発中です。")

    def update_timer_display(self, is_initial=False):
        if not is_initial:
            self.remaining_time -= 1
            if self.remaining_time <= 0:
                self.remaining_time = 0
                if self.timer.isActive():
                    self.timer.stop()
                self.timer_paused = False
                QMessageBox.information(self, "試験終了", "試験時間が終了しました。")
        if self.timer.isActive():
            self.timer_paused = False
        hours, rem = divmod(self.remaining_time, 3600); mins, secs = divmod(rem, 60)
        self.timer_label.setText(f"{hours:02}:{mins:02}:{secs:02}")
        
    def toggle_memo_window(self):
        if not self.memo_window:
            self.memo_window = MemoWindow(self)
        if self.memo_window.isVisible():
            self.memo_window.hide()
        else:
            self.memo_window.show_overlay(self.memo_window.mode)

    def clear_highlight(self, text_edit):
        cursor = text_edit.textCursor(); cursor.select(QTextCursor.SelectionType.Document)
        default_format = QTextCharFormat(); cursor.setCharFormat(default_format)
        cursor.clearSelection(); text_edit.setTextCursor(cursor)
        
    def highlight_all_in_answer(self, sheet, keyword):
        for editor in sheet.pages:
            self.clear_highlight(editor)
            if not keyword:
                continue
            highlight_format = QTextCharFormat()
            highlight_format.setBackground(QColor("yellow"))
            doc = editor.document()
            cursor = QTextCursor(doc)
            while True:
                cursor = doc.find(keyword, cursor)
                if cursor.isNull():
                    break
                cursor.mergeCharFormat(highlight_format)
        sheet.search_count_label.setText(self._count_keyword_occurrences(sheet, keyword))

    def replace_in_answer(self, sheet):
        keyword = sheet.answer_search_input.text()
        replace_text = sheet.answer_replace_input.text()
        editor = sheet.current_editor()
        cursor = editor.textCursor()
        if keyword and cursor.hasSelection() and cursor.selectedText() == keyword:
            cursor.insertText(replace_text)
        editor.find(keyword)
        sheet.update_status_label()
        self.update_char_count()
        sheet.contentChanged.emit()

    def find_in_answer(self, sheet, forward=True):
        keyword = sheet.answer_search_input.text()
        if not keyword:
            return
        editor = sheet.current_editor()
        flags = QTextDocument.FindFlag(0)
        if not forward:
            flags |= QTextDocument.FindFlag.FindBackward
        editor.find(keyword, flags)

    def replace_all_in_answer(self, sheet):
        keyword = sheet.answer_search_input.text()
        replace_text = sheet.answer_replace_input.text()
        if not keyword:
            return
        for editor in sheet.pages:
            self.clear_highlight(editor)
            text = editor.toPlainText()
            editor._internal_change = True
            editor.setPlainText(text.replace(keyword, replace_text))
            editor._internal_change = False
        sheet.update_status_label()
        self.update_char_count()
        sheet.contentChanged.emit()

    def _count_keyword_occurrences(self, sheet, keyword):
        if not keyword:
            return "0/0"
        total = 0
        for editor in sheet.pages:
            text = editor.toPlainText()
            total += text.count(keyword)
        return f"0/{total}" if total else "0/0"

    def update_char_count(self):
        for i in range(self.answer_tab_widget.count()):
            sheet = self.answer_tab_widget.widget(i)
            sheet.update_status_label()

    def open_manual(self): QDesktopServices.openUrl(QUrl.fromLocalFile("manual.pdf"))
    
    def open_pdf_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "PDF", "", "*.pdf")
        if file_path:
            file_name = os.path.basename(file_path)
            self.exam_type_label.setText(file_name)
            
            if self.page_overview_dialog:
                self.page_overview_dialog.close()
            self.clear_selection()
            self.reset_zoom()
            if self.spread_toggle_action.isChecked():
                self.spread_toggle_action.blockSignals(True)
                self.spread_toggle_action.setChecked(False)
                self.spread_toggle_action.blockSignals(False)
            self.spread_mode = False
            self.scroll_toggle_action.blockSignals(True)
            self.scroll_toggle_action.setChecked(False)
            self.scroll_toggle_action.blockSignals(False)
            self._apply_scroll_settings_without_refresh(False)
            for i in range(self.answer_tab_widget.count()):
                sheet = self.answer_tab_widget.widget(i)
                sheet.set_page_texts([''] * sheet.TOTAL_PAGES)
                sheet.set_current_page(0)
            self._clear_text_annotations()
            self._clear_shape_annotations()
            self.current_pdf_path = file_path
            self.pdf_document = fitz.open(file_path); self.total_pages = len(self.pdf_document); self.current_page = 0; self.show_page(self.current_page)
            self.undo_stack.clear()
            self.redo_stack.clear()
            self.register_snapshot()
            self.update_char_count()
        
    def show_page(self, page_number):
        if not self.pdf_document or not (0 <= page_number < self.total_pages): return
        self.current_page = page_number
        self._update_text_annotations_visibility()
        self._update_shape_annotations_visibility()
        page = self.pdf_document.load_page(page_number)
        page_rect = page.rect

        viewport = self.pdf_scroll_area.viewport() if self.pdf_scroll_area else None
        if viewport:
            label_width = viewport.width()
            label_height = viewport.height()
        else:
            target_rect = self.pdf_display_label.contentsRect()
            label_width = target_rect.width()
            label_height = target_rect.height()
        if label_width <= 0 or label_height <= 0:
            self._pending_page_render = page_number
            if not self._render_timer_active:
                self._render_timer_active = True
                QTimer.singleShot(0, self._retry_show_page)
            return

        if self.spread_mode and not self.scroll_toggle_action.isChecked():
            effective_width = max(1, label_width - 20)
            width_per_page = effective_width / 2
            scale_x = width_per_page / page_rect.width if page_rect.width else 1.0
        else:
            scale_x = label_width / page_rect.width if page_rect.width else 1.0
        scale_y = label_height / page_rect.height if page_rect.height else 1.0
        fit_scale = max(min(scale_x, scale_y), 0.1) * self.zoom_factor

        window_handle = self.windowHandle()
        dpr = float(window_handle.devicePixelRatio()) if window_handle else 1.0

        primary_image = self._render_page_image(page, fit_scale, dpr)
        if self.spread_mode and not self.scroll_toggle_action.isChecked():
            if page_number + 1 < self.total_pages:
                right_page = self.pdf_document.load_page(page_number + 1)
                right_image = self._render_page_image(right_page, fit_scale, dpr)
            else:
                right_image = None
            combined_image = self._compose_spread_image(primary_image, right_image, dpr)
            pixmap = QPixmap.fromImage(combined_image)
        else:
            pixmap = QPixmap.fromImage(primary_image)
        pixmap.setDevicePixelRatio(dpr)

        self.pdf_display_label.setPixmap(pixmap)
        self.pdf_display_label.adjustSize()
        self.pdf_display_label._clamp_all_annotations()
        self.page_label.setText(f"{self.current_page + 1} / {self.total_pages}")
        self.page_num_input.setText(str(self.current_page + 1))
        self._pending_page_render = None
        self._update_text_annotations_visibility()
        self._update_shape_annotations_visibility()

        if self.pdf_scroll_area:
            hbar = self.pdf_scroll_area.horizontalScrollBar()
            vbar = self.pdf_scroll_area.verticalScrollBar()
            h_offset = max(0, (self.pdf_display_label.width() - self.pdf_scroll_area.viewport().width()) // 2)
            v_offset = max(0, (self.pdf_display_label.height() - self.pdf_scroll_area.viewport().height()) // 2)
            hbar.setValue(h_offset)
            vbar.setValue(v_offset)
        
    def show_prev_page(self): self.show_page(self.current_page - 1)
    
    def show_next_page(self): self.show_page(self.current_page + 1)
    
    def resizeEvent(self, event):
        super().resizeEvent(event)
        if hasattr(self, 'pdf_document') and self.pdf_document:
            self._pending_page_render = self.current_page
            if not self._render_timer_active:
                self._render_timer_active = True
                QTimer.singleShot(0, self._retry_show_page)

    def _retry_show_page(self):
        self._render_timer_active = False
        if self._pending_page_render is not None:
            page_number = self._pending_page_render
            self._pending_page_render = None
            self.show_page(page_number)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
