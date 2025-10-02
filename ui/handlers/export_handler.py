from __future__ import annotations
import os
import json
from datetime import datetime
from typing import TYPE_CHECKING, Dict, List

import fitz
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PyQt6.QtWidgets import QFileDialog, QMessageBox
from PyQt6.QtGui import QFont, QFontInfo, QPainter, QImage
from PyQt6.QtCore import QMarginsF, QPointF, QRectF, Qt, QPageLayout
from PyQt6.QtPrintSupport import QPrinter

from ui.widgets import ANSWER_TEMPLATE_PATH

if TYPE_CHECKING:
    from ..main_window import MainWindow
    from ..widgets.answer_sheet_widget import AnswerSheetWidget

class ExportHandler:
    """
    答案データやセッション情報を外部ファイル形式（Word, PDF, JSON）にエクスポートする機能を提供します。
    """
    def __init__(self, main_window: MainWindow) -> None:
        """
        ExportHandlerのコンストラクタ。

        Args:
            main_window (MainWindow): 親となるメインウィンドウインスタンス。
        """
        self.main: MainWindow = main_window
        self._template_image_cache: Dict[int, QImage] = {}

    def save_as_word(self) -> None:
        """
        現在アクティブな答案をWord (.docx) 形式で保存する。
        ファイルダイアログを表示し、ユーザーに出力先を指定させます。
        """
        current_sheet = self.main.answer_tab_widget.currentWidget()
        if not current_sheet:
            return

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        default_name = f"answer_{timestamp}.docx"
        initial_path = os.path.join(os.path.expanduser("~"), default_name)

        file_path, _ = QFileDialog.getSaveFileName(
            self.main, "Word形式で保存", initial_path, "Word Documents (*.docx)"
        )
        if not file_path:
            return
        if not file_path.lower().endswith('.docx'):
            file_path += '.docx'

        try:
            self._export_sheet_to_docx(current_sheet, file_path)
            QMessageBox.information(self.main, "保存完了", f"答案をWord形式で保存しました。\n{file_path}")
        except Exception as exc:
            QMessageBox.critical(self.main, "保存エラー", f"Wordファイルの保存に失敗しました。\n{exc}")

    def save_as_pdf(self) -> None:
        """
        現在アクティブな答案をPDF形式で保存する。
        ファイルダイアログを表示し、ユーザーに出力先を指定させます。
        """
        current_sheet = self.main.answer_tab_widget.currentWidget()
        if not current_sheet:
            return

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        default_name = f"answer_{timestamp}.pdf"
        initial_path = os.path.join(os.path.expanduser("~"), default_name)

        file_path, _ = QFileDialog.getSaveFileName(
            self.main, "PDF形式で保存", initial_path, "PDF Files (*.pdf)"
        )
        if not file_path:
            return
        if not file_path.lower().endswith('.pdf'):
            file_path += '.pdf'

        try:
            self._export_sheet_to_pdf(current_sheet, file_path)
            QMessageBox.information(self.main, "保存完了", f"答案をPDF形式で保存しました。\n{file_path}")
        except Exception as exc:
            QMessageBox.critical(self.main, "保存エラー", f"PDFファイルの保存に失敗しました。\n{exc}")

    def _export_sheet_to_docx(self, sheet: AnswerSheetWidget, file_path: str) -> None:
        """答案シートの内容をWord文書として書き出す内部メソッド。"""
        page_texts = self._sheet_page_texts(sheet)
        doc = Document()
        section = doc.sections[0]
        section.page_width, section.page_height = Cm(21.0), Cm(29.7) # A4
        section.left_margin, section.right_margin = Cm(2.0), Cm(2.0)
        section.top_margin, section.bottom_margin = Cm(2.0), Cm(2.0)

        style = doc.styles['Normal']
        style.font.name = 'MS Mincho'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')

        for page_index, page_text in enumerate(page_texts):
            lines = self._page_lines_for_export(page_text)
            table = doc.add_table(rows=23, cols=2)
            table.style = 'Table Grid'
            table.autofit = False

            line_col, text_col = table.columns[0], table.columns[1]
            line_width = Cm(1.0)
            line_col.width = line_width
            text_col.width = section.page_width - section.left_margin - section.right_margin - line_width

            for row_idx, row in enumerate(table.rows):
                line_cell, text_cell = row.cells[0], row.cells[1]
                line_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                line_cell.paragraphs[0].add_run(f"{row_idx + 1:02}")
                text_cell.paragraphs[0].add_run(lines[row_idx] if row_idx < len(lines) else '')

            if page_index < len(page_texts) - 1:
                doc.add_page_break()
        doc.save(file_path)

    def _export_sheet_to_pdf(self, sheet: AnswerSheetWidget, file_path: str) -> None:
        """答案シートの内容をPDF文書として書き出す内部メソッド。"""
        page_texts = self._sheet_page_texts(sheet)
        if not page_texts: page_texts = ['']

        font = QFont("Hiragino Mincho ProN", 12)
        if not QFontInfo(font).exactMatch(): font = QFont("MS Mincho", 12)

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(file_path)
        printer.setPageSize(QPrinter.PageSize.A4)
        printer.setPageMargins(QMarginsF(20, 20, 20, 25), QPageLayout.Unit.Millimeter)

        painter = QPainter(printer)
        painter.setFont(font)
        metrics = painter.fontMetrics()
        line_height = metrics.lineSpacing()
        page_rect = printer.pageRect(QPrinter.Unit.Point)
        template_image = self._get_template_image(printer.resolution())

        max_lines, max_chars = (sheet.pages[0].max_lines, sheet.pages[0].max_chars) if sheet.pages else (23, 30)

        for i, page_text in enumerate(page_texts):
            if i > 0: printer.newPage()
            if not template_image.isNull():
                painter.drawImage(page_rect, template_image)

            lines = self._page_lines_for_export(page_text, max_lines, max_chars)
            for row, line in enumerate(lines):
                # NOTE: 座標はテンプレートに合わせて調整が必要
                painter.drawText(QPointF(100, 100 + row * line_height), line)

            painter.drawText(QRectF(page_rect), Qt.AlignmentFlag.AlignBottom | Qt.AlignmentFlag.AlignHCenter, f"{i + 1} / {len(page_texts)}")
        painter.end()

    def _sheet_page_texts(self, sheet: AnswerSheetWidget) -> List[str]:
        """答案シートから空でない全ページのテキストを取得する。"""
        texts = sheet.get_page_texts()
        last_nonempty = -1
        for idx, text in enumerate(texts):
            if text.strip():
                last_nonempty = idx
        # 最後に入力があったページまでを返す。全く入力がなければ最初のページのみ返す
        return texts[:last_nonempty + 1] if last_nonempty != -1 else ([''] if not texts else [texts[0]])

    def _page_lines_for_export(self, text: str, max_lines: int = 23, max_chars: int = 30) -> List[str]:
        """1ページ分のテキストを、指定された最大行数と文字数で折り返して行リストを生成する。"""
        lines = []
        for raw_line in (text.split('\n') if text else ['']):
            for i in range(0, len(raw_line), max_chars):
                lines.append(raw_line[i:i+max_chars])
        return lines[:max_lines]

    def _get_template_image(self, dpi: int) -> QImage:
        """
        指定されたDPIでレンダリングされた答案テンプレート画像を返す。
        パフォーマンス向上のため、一度レンダリングした画像はキャッシュする。
        """
        if dpi in self._template_image_cache:
            return self._template_image_cache[dpi]

        image = QImage()
        if ANSWER_TEMPLATE_PATH and os.path.exists(ANSWER_TEMPLATE_PATH):
            try:
                doc = fitz.open(ANSWER_TEMPLATE_PATH)
                page = doc.load_page(0)
                mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                image.loadFromData(pix.tobytes("png"), "PNG")
            except Exception:
                pass # テンプレートの読み込みに失敗しても処理は続行

        self._template_image_cache[dpi] = image
        return image

    def save_session_via_dialog(self) -> bool:
        """
        現在のセッション全体の状態をJSONファイルとして保存する。
        ファイルダイアログを表示し、ユーザーに出力先を指定させる。

        Returns:
            bool: 保存が成功した場合はTrue、キャンセルまたは失敗した場合はFalse。
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        default_name = f"cbt_session_{timestamp}.json"
        initial_path = os.path.join(os.path.expanduser("~"), default_name)

        file_path, _ = QFileDialog.getSaveFileName(
            self.main, "セッションの保存", initial_path, "CBTセッション (*.json);;All Files (*)"
        )
        if not file_path:
            return False

        try:
            data = self.main.history_handler.serialize_full_session()
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            QMessageBox.information(self.main, "保存完了", f"セッションを保存しました。\n{file_path}")
            return True
        except Exception as exc:
            QMessageBox.critical(self.main, "保存エラー", f"セッションの保存に失敗しました。\n{exc}")
            return False